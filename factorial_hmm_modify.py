# -*- coding: utf-8 -*-
"""
Created on Fri Mar 13 10:42:59 2020

@author: lenovo
"""
import hmmlearn.hmm as hmm
import numpy as np
import itertools
import pandas as pd
from scipy.stats import norm
from scipy.optimize import minimize
from docx import Document
import matplotlib.pyplot as plt
import matplotlib as mpl
import matplotlib.dates as mdates
from datetime import datetime
import time

def read_clusters(file_path):
    doc=Document(file_path)
    cluster_num=len(doc.tables[3].rows)-1
    clusters=[]
    for i in range(1,len(doc.tables[3].rows)):
        clusters.append(float(doc.tables[3].cell(i,1).text))
    return (cluster_num,clusters)
 
    
def read_dataset(file_path,require_time=None):
    f=open(file_path,'r')
    value=[]
    time_line=[]
    line=f.readline()
    while line:
        data=line.split(' ')
        value.append(int(data[1][0:len(data[1])-1]))
        time_line.append(int(data[0]))
        line=f.readline()
    f.close()
    if require_time==None:
        return value
    else:
        return value,time_line
    
       

class factorial_hmm(object):
    def __init__(self,HMMs_num,house):
        self.HMMs_num=HMMs_num   #隐马尔科夫链数目，即总电器数量
        self.house=house         #数据集对应房间序号
        self.HMMs_para={}        #存放各条HMM参数的字典
        self.states_combination=[]  #各条hmm状态组合，形如('1','2','3','4',...)该tuple长度为hmm隐状态数量
        self.startprob_=None      #fhmm初始状态分布
        self.transmat_=None       #fhmm状态转移矩阵
        self.means_=None          #fhmm观测状态概率密度均值矩阵
        self.covars_=None         #fhmm观测状态概率密度协方差矩阵
        self.n_components=None    #fhmm隐状态总数
        self.states_frame=None    #各条hmm隐状态识别表格
        self.observe_frame=None   #各条hmm负荷分解有功功率表格
        self.aggregate_frame=None #一段时间内合成的总的有功功率采样数据表格
        self.true_value_frame_dic={} #一段时间内各个电器实际的有功采样数据表格组成的字典，索引为channel
         
    def get_HMMs_para(self):
        '''
        利用每个电器的有功采样数据和之前聚类得到的每个电器的状态数，
        学习得出每条hmm参数，结果存放在self.HMMs_para字典中，key为channel，value为list，
        依次存放pi、A、均值、协方差
        '''
        for j in range(2,self.HMMs_num+2):
            cluster_path='C:/Users/lenovo/Desktop/domestic UK2015/house'+str(self.house)+'/channel_'+str(j)+'.docx'
            states_num,means=read_clusters(cluster_path)
            dataset_path='C:/Users/lenovo/Desktop/house'+str(self.house)+'/channel_'+str(j)+'.dat'
            sample_value=read_dataset(dataset_path)
            HMM_model=hmm.GaussianHMM(n_components=states_num, covariance_type='full')
            HMM_model.fit(np.array([sample_value]).reshape(-1,1))
            self.HMMs_para['channel_'+str(j)]=[HMM_model.startprob_,HMM_model.transmat_,\
                           HMM_model.means_,HMM_model.covars_]
            
    def compute_FHMM_para(self):
        '''
        调用计算各fhmm参数的函数，存放结果
        pi的列数，A的阶数，means、covars的元素数均为隐状态数即self.n_components
        所有参数的排列顺序均与变量self.states_combination内组合状态的排列顺序相对应一致
        '''
        self.startprob_=self.compute_pi_fhmm()
        self.transmat_=self.compute_A_fhmm()
        self.means_=self.compute_means_fhmm()
        self.covars_=self.compute_covars_fhmm()
        self.n_components=self.compute_n_components()
        
    def compute_pi_fhmm(self):
        '''
        利用克罗内科积计算fhmm的pi矩阵
        '''
        list_pi=[]
        for j in range(2,len(self.HMMs_para.keys())+2):
            list_pi.append(self.HMMs_para['channel_'+str(j)][0])
        result = list_pi[0]
        for i in range(len(list_pi) - 1):
            result = np.kron(result, list_pi[i + 1])
        return result
    
    def compute_A_fhmm(self):
        '''
        利用克罗内科积计算fhmm的A矩阵
        '''
        list_A=[]
        for j in range(2,len(self.HMMs_para.keys())+2):
            list_A.append(self.HMMs_para['channel_'+str(j)][1])
        result = list_A[0]
        for i in range(len(list_A) - 1):
            result = np.kron(result, list_A[i + 1])
        return result
    
    def compute_means_fhmm(self):
        '''
        利用itertools.product方法组合每个fhmm隐状态所对应每条hmm的均值，
        并根据每条hmm相互独立的性质将这些均值相加得到fhmm每个隐状态观测概率分布均值
        
        根据每个hmm的均值列表中均值大小排列方式，生成与这些大小排列方式相同的从'1'开始的数字字符串，
        即均值最小的位置对应字符串位置上的数字为'1'，依此类推，
        然后利用itertools.product方法对每个均值列表对应的字符串进行组合，
        得到states_combination，这一list变量的长度为fhmm隐状态数，每一个元素代表了一种fhmm的隐状态，
        每一个元素为多个单数字字符组成的tuple，tuple长度为hmm条数，代表了该fhmm隐状态由每条hmm链的第几大小的均值组成的
        states_combination中元素排列顺序与pi、A、means、covars元素顺序都是相对应的
        '''
        list_means=[]
        result=[]
        for j in range(2,len(self.HMMs_para.keys())+2):
            list_means.append(self.HMMs_para['channel_'+str(j)][2].reshape(1,-1)[0])
        for i in itertools.product(*list_means):
            result.append(sum(i))
        states_number=[]
        for k in list_means:
            string=''
            sort_means=sorted(k)
            for p in k:
                for t in range(len(sort_means)):
                    if p==sort_means[t]:
                        string+=str(t+1)
                        break
            states_number.append(string)
        for q in itertools.product(*states_number):
            self.states_combination.append(q)
        return np.array([result]).reshape(-1,1)
    
    def compute_covars_fhmm(self):
        '''
        计算fhmm协方差矩阵
        '''
        list_covars=[]
        result=[]
        for j in range(2,len(self.HMMs_para.keys())+2):
            list_covars.append(self.HMMs_para['channel_'+str(j)][3].reshape(1,-1)[0])
        for i in itertools.product(*list_covars):
            result.append(sum(i))
        return np.array([result]).reshape(-1,1,1)

    def compute_n_components(self):
        '''
        计算fhmm隐状态数
        '''
        return len(self.states_combination)
    
    def get_aggregate_frame(self,time_start,time_len):
        '''
        获得time_range时间范围内所有电器实际有功功率数值的加和，
        返回一个frame，index为该段时间的时间点，values为实际有功加和，
        结果作为之后负荷功率分解的总有功约束条件，和用来预测fhmm的隐状态
        '''
        for i in range(2,self.HMMs_num+2):
            dataset_path='C:/Users/lenovo/Desktop/house'+str(self.house)+'/channel_'+str(i)+'.dat'
            sample_value,sample_time=read_dataset(file_path=dataset_path,require_time='yes')
            time_frame=self.get_a_time_frame(time_start,time_len,sample_time,sample_value)
            if i!=2:
                time_frame.index=self.true_value_frame_dic['channel_2'].index
            self.true_value_frame_dic['channel_'+str(i)]=time_frame
    
        aggregate_value=[]
        for j in range(int(time_len)):
            value=0
            for k in range(2,self.HMMs_num+2):
                value+=self.true_value_frame_dic['channel_'+str(k)].values[j,0]
            aggregate_value.append(value)
        self.aggregate_frame=pd.DataFrame(np.array([aggregate_value]).reshape(-1,1),\
                                          index=list(self.true_value_frame_dic['channel_2'].index),columns=['aggregate_value'])
        
    
    def decode(self):
        '''
        利用之前计算得到的一段时间内的总有功数值和合成的fhmm参数，预测得到fhmm隐状态序列
        然后根据对应关系得到self.states_combination元素的排列，分解得到每一时刻下各hmm的隐状态
        返回以该段时间点为index，channel为columns的状态frame
        '''
        fhmm_model=hmm.GaussianHMM(n_components=self.n_components, covariance_type='full')
        fhmm_model.startprob_=self.startprob_
        fhmm_model.transmat_=self.transmat_
        fhmm_model.means_=self.means_
        fhmm_model.covars_=self.covars_
        logprob,state_sequence=fhmm_model.decode(self.aggregate_frame.values)
        state_sequence_str={}
        for i in range(2,self.HMMs_num+2):
            state_str=[]
            for j in state_sequence:
                state_str.append(int(self.states_combination[j][i-2]))
            state_sequence_str['channel_'+str(i)]=state_str
        self.states_frame=pd.DataFrame(state_sequence_str,index=list(self.aggregate_frame.index))
        
        
    def get_a_time_frame(self,time_start,time_len,full_time_list,full_value_list):
        '''
        根据时间范围截取时间点和对应的有功采样数值，返回一个frame
        '''
        time_list=[]
        for time in full_time_list:
            if time>=time_start:
                time_list.append(time)
            if len(time_list)==time_len:
                break
        value_list=full_value_list[full_time_list.index(time_list[0]):\
                                   full_time_list.index(time_list[-1])+1]
        time_frame=pd.DataFrame({'ture_value':value_list},index=time_list)
        return time_frame
    
    def get_observe_frame(self):
        '''
        有功负荷分解
        根据之前获得到的每条hmm链每个时刻的隐状态，计算一个条件极值
        条件为总的有功约束即之前合成得到的，极值为所有电器在各自的某一隐状态下观测有功功率概率乘积最大化
        '''
        self.observe_frame=pd.DataFrame(np.empty((len(self.aggregate_frame.index),self.HMMs_num)),\
                                        index=list(self.aggregate_frame.index),\
                                        columns=list(self.states_frame.columns))
        for i in list(self.aggregate_frame.index):
            guass_func_list=[]
            x0=[]
            e=1e-10
            for j in range(2,self.HMMs_num+2):
                state=int(self.states_frame.loc[i,'channel_'+str(j)])
                meanlist=list(self.HMMs_para['channel_'+str(j)][2].reshape(1,-1)[0])
                mean=sorted(meanlist)[state-1]
                covar=self.HMMs_para['channel_'+str(j)][3].reshape(1,-1)[0][meanlist.index(mean)]
                guass_func_list.append(norm(loc=mean, scale=covar**0.5))
                x0.append(mean)
            fun=lambda o:(guass_func_list[0].logpdf(o[0])+guass_func_list[1].logpdf(o[1])
                         +guass_func_list[2].logpdf(o[2])
                         +guass_func_list[3].logpdf(o[3])
                         +guass_func_list[4].logpdf(o[4])
                         #+guass_func_list[5].logpdf(o[5])
                         )**-1
            cons=({'type':'eq','fun':lambda o: o[0]+o[1]\
                   +o[2]\
                   +o[3]\
                   +o[4]\
                   #+o[5]\
                   -self.aggregate_frame.loc[i,'aggregate_value']},
                   {'type':'ineq','fun':lambda o:o[0]-e},
                   {'type':'ineq','fun':lambda o:o[1]-e},
                   {'type':'ineq','fun':lambda o:o[2]-e},
                   {'type':'ineq','fun':lambda o:o[3]-e},
                   {'type':'ineq','fun':lambda o:o[4]-e},
                   #{'type':'ineq','fun':lambda o:o[5]-e}
                   )
            res=minimize(fun, x0, method='SLSQP', constraints=cons)
            for k in range(2,self.HMMs_num+2):
                self.observe_frame.loc[i,'channel_'+str(k)]=res.x[k-2]
            
            
    
                
    def compute_F1_measure(self):
        '''
        计算评价负荷分解精度指标F1_measure，越接近于1说明分解精度越高
        '''
        pre_list=[]
        rec_list=[]
        F1_rspectively=[]
        for i in range(2,self.HMMs_num+2):
            tp=[]
            ture_values=[]
            predict_values=[]
            for j in range(len(self.observe_frame.index)):
                tp.append(min(self.observe_frame.values[j,i-2],\
                              self.true_value_frame_dic['channel_'+str(i)].values[j,0]))
                ture_values.append(self.true_value_frame_dic['channel_'+str(i)].values[j,0])
                predict_values.append(self.observe_frame.values[j,i-2])
            pre_i=sum(tp)/sum(predict_values)
            if set(ture_values)==set([0]):
                rec_i=1
            else:
                rec_i=sum(tp)/sum(ture_values)
            f1_i=2*pre_i*rec_i/(pre_i+rec_i)
            F1_rspectively.append(f1_i)
            pre_list.append(pre_i)
            rec_list.append(rec_i)
        pre=np.mean(pre_list)
        rec=np.mean(rec_list)
        F1_general=2*pre*rec/(pre+rec)
        return F1_general,F1_rspectively
        
    
    
    def compute_NDE(self):
        '''
        计算评价负荷分解精度指标Normalized disaggregation error(NDE)，越接近于0分解精度越高
        '''
        subtract_squre_i=[]
        predict_squre_i=[]
        NDE_respectively=[]
        for i in range(2,self.HMMs_num+2):
            subtract_squre_t=[]
            predict_squre_t=[]
            for j in range(len(self.observe_frame.index)):
                subtract_squre_t.append((self.observe_frame.values[j,i-2]-\
                                           self.true_value_frame_dic['channel_'+str(i)].values[j,0])**2)
                predict_squre_t.append((self.observe_frame.values[j,i-2])**2)
            subtract_squre_t_sum=sum(subtract_squre_t)
            predict_squre_t_sum=sum(predict_squre_t)
            NDE_respectively.append(np.sqrt(subtract_squre_t_sum/predict_squre_t_sum))
            subtract_squre_i.append(subtract_squre_t_sum)
            predict_squre_i.append(predict_squre_t_sum)
        NDE_general=np.sqrt(sum(subtract_squre_i)/sum(predict_squre_i))
        return NDE_general,NDE_respectively
        
    def energy_consumption_portion(self):
        aggregate_consumption=0
        true_indiv_consumption_protion=[]
        observe_indiv_consumption_protion=[]
        for i in range(len(self.aggregate_frame.index)):
            if i<len(self.aggregate_frame.index)-1:
                aggregate_consumption+=self.aggregate_frame.values[i,0]*(self.aggregate_frame.index[i+1]-\
                                                              self.aggregate_frame.index[i])
        for j in range(2,self.HMMs_num+2):
            true_indiv_aggregate=0
            observe_indiv_aggregate=0
            for k in range(len(self.aggregate_frame.index)):
                if k<len(self.aggregate_frame.index)-1:
                    true_indiv_aggregate+=self.true_value_frame_dic['channel_'+str(j)].values[k,0]*\
                                          (self.true_value_frame_dic['channel_'+str(j)].index[k+1]-\
                                           self.true_value_frame_dic['channel_'+str(j)].index[k])
                    observe_indiv_aggregate+=self.observe_frame.values[k,j-2]*\
                                            (self.observe_frame.index[k+1]-\
                                             self.observe_frame.index[k])
            true_indiv_consumption_protion.append(true_indiv_aggregate/aggregate_consumption) 
            observe_indiv_consumption_protion.append(observe_indiv_aggregate/aggregate_consumption)
        return true_indiv_consumption_protion,observe_indiv_consumption_protion
       
    
    def compute_MAE(self):
        err_list=[]
        MAE_list=[]
        for i in range(2,self.HMMs_num+2):
            for j in range(len(self.observe_frame.index)):
                err_list.append(abs(self.true_value_frame_dic['channel_'+str(i)].values[j,0]-\
                                                          self.observe_frame.values[j,i-2]))
            MAE_list.append(np.mean(err_list[-len(self.observe_frame.index):]))
        return np.mean(err_list),MAE_list
    
    def compute_RMSE(self):
        err_squre_list=[]
        RMSE_list=[]
        for i in range(2,self.HMMs_num+2):
            for j in range(len(self.observe_frame.index)):
                err_squre_list.append((self.true_value_frame_dic['channel_'+str(i)].values[j,0]-\
                                                          self.observe_frame.values[j,i-2])**2)
            RMSE_list.append((np.mean(err_squre_list[-len(self.observe_frame.index):]))**0.5)
        return (np.mean(err_squre_list))**0.5,RMSE_list
        
            

'''
使用该类时，先根据HMMs_num的数值修改条件极值目标函数和约束项中的变量个数

'''  

house3_fhmm=factorial_hmm(HMMs_num=5,house=6)
house3_fhmm.get_HMMs_para()
house3_fhmm.compute_FHMM_para()
house3_fhmm.get_aggregate_frame(time_start=1.40577e9,time_len=5e3)
house3_fhmm.decode()
house3_fhmm.get_observe_frame()
print(house3_fhmm.aggregate_frame)
print(house3_fhmm.states_frame)
print(house3_fhmm.observe_frame)
print(house3_fhmm.n_components)


for i in range(2,house3_fhmm.HMMs_num+2):
    GMT_line=[]
    time_num=list(house3_fhmm.observe_frame.index)
    for j in time_num:
        GMT_stamp=int(j)-7*3600
        GMT_time=time.localtime(GMT_stamp)
        GMT=time.strftime('%m%d%H%M%S',GMT_time)
        GMT_line.append(GMT)
    
    x=[datetime.strptime(d,'%m%d%H%M%S') for d in GMT_line]
    
    fig=plt.figure()
    ax1=fig.add_subplot(1,1,1)
    ax1.plot(x,list(house3_fhmm.observe_frame['channel_'+str(i)]),label='disaggragate')
    plt.gcf().autofmt_xdate()
    
    allhours=mdates.HourLocator()
    ax1.xaxis.set_major_locator(allhours) 
    ax1.xaxis.set_major_formatter(mdates.DateFormatter('%m%d%H'))
    
    minutesLoc=mpl.dates.MinuteLocator(interval=30)
    ax1.xaxis.set_minor_locator(minutesLoc)
    ax1.xaxis.set_minor_formatter(mdates.DateFormatter('%M'))
    
    ax1.tick_params(pad=15)
    
    mpl.rcParams['xtick.direction'] = 'in'
    mpl.rcParams['ytick.direction'] = 'in'
    
    ax1.set_xlabel('date(month/day/hour/minute)')
    ax1.set_ylabel('active power(W)')
    
    plt.legend()
    plt.show()


    fig=plt.figure()
    ax1=fig.add_subplot(1,1,1)
    ax1.plot(x,list(house3_fhmm.true_value_frame_dic['channel_'+str(i)]['ture_value']),label='real_data')
    plt.gcf().autofmt_xdate()
    
    allhours=mdates.HourLocator()
    ax1.xaxis.set_major_locator(allhours) 
    ax1.xaxis.set_major_formatter(mdates.DateFormatter('%m%d%H'))
    
    minutesLoc=mpl.dates.MinuteLocator(interval=30)
    ax1.xaxis.set_minor_locator(minutesLoc)
    ax1.xaxis.set_minor_formatter(mdates.DateFormatter('%M'))
    
    ax1.tick_params(pad=15)
    
    mpl.rcParams['xtick.direction'] = 'in'
    mpl.rcParams['ytick.direction'] = 'in'
    
    ax1.set_xlabel('date(month/day/hour/minute)')
    ax1.set_ylabel('active power(W)')
    
    plt.legend()
    plt.show()
    
print(house3_fhmm.compute_F1_measure())
print(house3_fhmm.compute_NDE())
print(house3_fhmm.energy_consumption_portion())


'''
for i in range(2,7):
    dataset_path='C:/Users/lenovo/Desktop/house6/channel_'+str(i)+'.dat'
    value,time_line=read_dataset(dataset_path,require_time=1)
    plt.figure()
    plt.plot(time_line,value)
    plt.xlim(1.40577e9,1.40580e9)
    plt.show()

'''

'''
house6's (time_start=1.40577e9,time_len=5e3)
F1,NDE,energy_consumption_portion results:
(0.7740271489836756, [0.8698686861966736, 0.4464345823221753, 0.9110389337704555,
                      0.32384797489430045, 0.9492244861319984])
(0.6463040826521173, [0.43282973154523857, 1.5732911621786192, 0.38001128921546207, 
                      0.8958576607694799, 0.07488400019851343])
([0.12705218541841437, 0.1625255425617239, 0.14574949893684214, 0.03308794529271243, 0.5315848277903071], 
[0.1507143248705586, 0.058443302675984854, 0.130119301804774, 0.17377834338754372, 0.4869447272611321])

'''

'''
plt.figure()
app=['overall','fridge','oven','elec_hob','kettle','dishwasher']
f1_list=np.array([0.7740271489836756, 0.8698686861966736, 0.4464345823221753, 0.9110389337704555,
                  0.32384797489430045, 0.9492244861319984])*100
plt.bar(np.linspace(0,1,6),f1_list,tick_label=app,color=['b', 'c', 'm', 'y','r','g'],width=0.1)
plt.xlim(-0.2,1.2)
plt.ylim(0,100)
plt.ylabel('F1(100%)')
plt.show()


plt.figure()
app=['overall','fridge','oven','elec_hob','kettle','dishwasher']
f1_list=np.array([0.6463040826521173, 0.43282973154523857, 1.5732911621786192, 0.38001128921546207, 
                  0.8958576607694799, 0.07488400019851343])
plt.bar(np.linspace(0,1,6),f1_list,tick_label=app,color=['b', 'c', 'm', 'y','r','g'],width=0.1)
plt.xlim(-0.2,1.2)
plt.ylim(0,1.8)
plt.ylabel('TERR')
plt.show()


plt.figure()
algri=['true','FHMM']
fridge_freezer=np.array([0.12705218541841437,0.1507143248705586])*100
oven=np.array([0.1625255425617239,0.058443302675984854])*100
electric_hob=np.array([0.14574949893684214,0.130119301804774])*100
kettle=np.array([0.03308794529271243,0.17377834338754372])*100
dishwasher=np.array([0.5315848277903071,0.4869447272611321])*100
plt.bar((0,0.5),fridge_freezer,label='fridge',fc='c',width=0.25)
plt.bar((0,0.5),oven,label='oven',fc='m',bottom=fridge_freezer,width=0.25)
plt.bar((0,0.5),electric_hob,label='elec_hob',fc='y',\
        bottom=np.array(fridge_freezer)+np.array(oven),tick_label=algri,width=0.25)
plt.bar((0,0.5),kettle,label='kettle',fc='r',\
        bottom=np.array(fridge_freezer)+np.array(oven)+np.array(electric_hob),width=0.25)
plt.bar((0,0.5),dishwasher,label='dishwasher',fc='g',\
        bottom=np.array(fridge_freezer)+np.array(oven)+np.array(electric_hob)+np.array(kettle),width=0.25)
plt.ylim(0,100)
plt.xlim(-0.5,1)
plt.legend(loc='upper right')
plt.ylabel('energy_consumption(100%)')
plt.show()
'''

'''
house6's (time_range=(1408000000-15*3600,1408000000-12*3600))
F1,NDE,energy_consumption_portion results:
(0.8548053446682986, [0.8276362699180588, 0.8823439537567355, 0.8279819777794757])
(0.4278197923280493, [0.48627852637503377, 0.3841387369246089, 0.4938501094480248])
([0.13303680164260334, 0.6387267687224214, 0.2282364296349753], 
 [0.175497838021514, 0.5394760455875082, 0.28502611639098546])
'''

'''
plt.figure()
app=['overall','fridge','oven','elec_hob']
f1_list=np.array([0.8548053446682986, 0.8276362699180588, 0.8823439537567355, 0.8279819777794757])*100
plt.bar(np.linspace(0,0.6,4),f1_list,tick_label=app,color=['b', 'c', 'm', 'y'],width=0.1)
plt.xlim(-0.2,0.8)
plt.ylim(0,100)
plt.ylabel('F1(100%)')
plt.show()

plt.figure()
app=['overall','fridge','oven','elec_hob']
f1_list=np.array([0.4278197923280493, 0.48627852637503377, 0.3841387369246089, 0.4938501094480248])
plt.bar(np.linspace(0,0.6,4),f1_list,tick_label=app,color=['b', 'c', 'm', 'y'],width=0.1)
plt.xlim(-0.2,0.8)
plt.ylim(0,1)
plt.ylabel('TERR')
plt.show()

plt.figure()
algri=['true','FHMM']
fridge_freezer=np.array([0.13303680164260334,0.175497838021514])*100
oven=np.array([0.6387267687224214,0.5394760455875082])*100
electric_hob=np.array([0.2282364296349753,0.28502611639098546])*100
plt.bar((0,0.5),fridge_freezer,label='fridge',fc='r',width=0.25)
plt.bar((0,0.5),oven,label='oven',fc='g',bottom=fridge_freezer,width=0.25)
plt.bar((0,0.5),electric_hob,label='elec_hob',fc='b',\
        bottom=np.array(fridge_freezer)+np.array(oven),tick_label=algri,width=0.25)
plt.ylim(0,100)
plt.xlim(-0.5,1)
plt.legend(loc='upper right')
plt.ylabel('energy_consumption(100%)')
plt.show()
'''

